"""Write a transcript as a formatted Word document (.docx)."""

from pathlib import Path


def write_docx(out_path, title, lines):
    """lines: [(start_s, end_s, text)] — speaker prefixes like
    'Tania: ...' are rendered with the name in bold."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(title, level=1)

    for start, _end, text in lines:
        para = doc.add_paragraph()
        stamp = para.add_run(f"[{int(start // 60):02d}:{int(start % 60):02d}] ")
        stamp.font.size = Pt(8)
        stamp.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        speaker, sep, rest = text.partition(": ")
        # treat the prefix as a speaker name only if it looks like one
        if sep and 0 < len(speaker) <= 30 and "\n" not in speaker:
            name_run = para.add_run(speaker + ": ")
            name_run.bold = True
            para.add_run(rest)
        else:
            para.add_run(text)

    doc.save(str(Path(out_path)))
    return Path(out_path)
